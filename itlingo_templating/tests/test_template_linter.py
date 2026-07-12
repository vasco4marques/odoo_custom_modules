import io
from unittest.mock import patch

from odoo.tests.common import BaseCase, tagged

from odoo.addons.itlingo_templating.services.template_linter import (
    extract_template_sources,
    lint_template,
)


REFERENCE = {
    "success": True,
    "types": [{
        "name": "Requirement",
        "properties": [{"name": "priority"}, {"name": "description"}],
    }],
    "profile_aliases": [{
        "alias": "requirements", "type_name": "Requirement",
    }],
    "root_alias": "project",
    "profile": {},
}


@tagged("post_install", "-at_install")
class TestTemplateLinter(BaseCase):

    def test_level_a_suggests_variable_and_type(self):
        result = lint_template(
            b"{{ requirments|length }} {{ by_type['Requiremnt']|length }}",
            "txt", REFERENCE,
        )

        self.assertEqual(len(result["findings"]), 2)
        suggestions = {item["suggestion"] for item in result["findings"]}
        self.assertEqual(suggestions, {"requirements", "Requirement"})
        self.assertEqual({item["category"] for item in result["findings"]}, {"dsl"})

    def test_typed_loop_field_is_checked_conservatively(self):
        result = lint_template(
            b"{% for req in requirements %}{{ req.priorty }}{% endfor %}",
            "md", REFERENCE,
        )
        self.assertEqual(len(result["findings"]), 1)
        self.assertEqual(result["findings"][0]["suggestion"], "priority")

        dynamic = lint_template(
            b"{% for req in collection %}{{ req.priorty }}{% endfor %}",
            "md", REFERENCE | {"profile": {"other_alias": "collection"}},
        )
        self.assertFalse(dynamic["findings"])

        reassigned = lint_template(
            b"{% for req in requirements %}{% set req = root %}{{ req.priorty }}{% endfor %}",
            "md", REFERENCE,
        )
        self.assertFalse(reassigned["findings"])

    def test_xlsx_extraction_reports_sheet_and_cells(self):
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Report"
        sheet["B2"] = "{{ requirements|length }}"
        output = io.BytesIO()
        workbook.save(output)

        extracted = extract_template_sources(output.getvalue(), "xlsx")
        self.assertEqual(extracted["sources"][0]["location"], "Report!B2")
        self.assertIn("requirements", extracted["sources"][0]["text"])

    def test_docx_patch_failure_degrades_to_top_level(self):
        template = type("Template", (), {
            "init_docx": lambda self: None,
            "get_xml": lambda self: "ignored",
            "patch_xml": lambda self, xml: (_ for _ in ()).throw(ValueError()),
            "get_undeclared_template_variables": lambda self, jinja_env: {"project"},
        })()
        with patch("docxtpl.DocxTemplate", return_value=template):
            extracted = extract_template_sources(b"not-a-docx", "docx")

        self.assertTrue(extracted["degraded"])
        self.assertEqual(extracted["top_level"], {"project"})

    def test_docx_extraction_uses_patched_xml(self):
        from docx import Document

        output = io.BytesIO()
        document = Document()
        document.add_paragraph("{{ project.title }}")
        document.save(output)

        extracted = extract_template_sources(output.getvalue(), "docx")
        self.assertFalse(extracted["degraded"])
        self.assertIn("project.title", extracted["sources"][0]["text"])

    def test_unavailable_inventory_still_checks_jinja(self):
        result = lint_template(
            b"{% for item in requirements %}{{ item.title }}",
            "txt", {"success": False},
        )
        self.assertFalse(result["skipped"])
        self.assertTrue(result["dsl_skipped"])
        self.assertEqual(result["findings"][0]["category"], "syntax")
        self.assertEqual(result["findings"][0]["location"], "line 1")

    def test_malformed_for_and_if_report_syntax_lines(self):
        for source, line in (
            (b"Header\n{% for item in requirements %}\n{{ item.title }}", "line 3"),
            (b"{% if project %}\n{{ project.title }}", "line 2"),
        ):
            result = lint_template(source, "md", REFERENCE)
            self.assertEqual(len(result["findings"]), 1)
            self.assertEqual(result["findings"][0]["category"], "syntax")
            self.assertEqual(result["findings"][0]["location"], line)

    def test_unknown_filter_and_test_have_suggestions_and_are_grouped(self):
        result = lint_template(
            b"{{ requirements | lenght }} {% if project is defiend %}ok{% endif %}",
            "txt", REFERENCE,
        )
        self.assertEqual(
            [(item["category"], item["suggestion"]) for item in result["findings"]],
            [("filter", "length"), ("test", "defined")],
        )
        self.assertEqual(
            [group["category"] for group in result["finding_groups"]],
            ["filter", "test"],
        )

    def test_valid_jinja_is_clean(self):
        result = lint_template(
            b"{% for item in requirements|sort(attribute='priority') %}"
            b"{% if item.priority is defined %}{{ item.priority|upper }}{% endif %}"
            b"{% endfor %}",
            "txt", REFERENCE,
        )
        self.assertFalse(result["findings"])

    def test_xlsx_column_a_convention_and_syntax_location(self):
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Report"
        sheet["B2"] = "{% for item in requirements %}"
        sheet["B3"] = "{{ item.title }}"
        output = io.BytesIO()
        workbook.save(output)

        result = lint_template(output.getvalue(), "xlsx", REFERENCE)
        categories = [item["category"] for item in result["findings"]]
        self.assertEqual(categories[0], "syntax")
        self.assertIn("structure", categories)
        self.assertTrue(all(
            item["location"].startswith("Report!")
            for item in result["findings"]
        ))

        clean_workbook = Workbook()
        clean_sheet = clean_workbook.active
        clean_sheet["A1"] = "{% for item in requirements %}"
        clean_sheet["A2"] = "{{ item.title }}"
        clean_sheet["A3"] = "{% endfor %}"
        clean_output = io.BytesIO()
        clean_workbook.save(clean_output)
        clean = lint_template(clean_output.getvalue(), "xlsx", REFERENCE)
        self.assertFalse(clean["findings"])

    def test_docx_unbalanced_table_row_markers_warn(self):
        from docx import Document

        output = io.BytesIO()
        document = Document()
        document.add_paragraph("{%tr for item in requirements %}")
        document.save(output)

        result = lint_template(output.getvalue(), "docx", REFERENCE)
        structure = [
            item for item in result["findings"] if item["category"] == "structure"
        ]
        self.assertEqual(len(structure), 1)
        self.assertIn("unbalanced", structure[0]["message"])

        clean_output = io.BytesIO()
        clean_document = Document()
        clean_document.add_paragraph("{%tr for item in requirements %}")
        clean_document.add_paragraph("{%tr endfor %}")
        clean_document.save(clean_output)
        clean_extracted = extract_template_sources(clean_output.getvalue(), "docx")
        self.assertFalse(clean_extracted["findings"])
